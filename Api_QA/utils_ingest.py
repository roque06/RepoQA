# utils_ingest.py
# ---------------------------------------------
# Extracción de texto desde adjuntos para usar como contexto en la generación
# - PDF:     PyMuPDF (texto nativo) + OCR por página con Tesseract si no hay texto
# - DOCX:    python-docx
# - TXT/CSV: decodificación estándar
# - XLSX:    pandas -> CSV
# - IMAGEN:  OCR con Tesseract
# ---------------------------------------------

import io
import csv
import hashlib
import platform
import shutil
from typing import List, Dict, Tuple

import pandas as pd
import fitz  # PyMuPDF

# ===== OCR (Tesseract) =====
try:
    from PIL import Image
    import pytesseract

    # Config multiplataforma: en Windows fijamos ruta; en Linux/Mac usamos PATH
    if platform.system() == "Windows":
        # Cambia esta ruta si instalaste Tesseract en otro lugar
        pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
    elif shutil.which("tesseract"):
        pytesseract.pytesseract.tesseract_cmd = shutil.which("tesseract")

    _OCR_OK = True
except Exception:
    _OCR_OK = False
    Image = None
    pytesseract = None

# ===== Extractores auxiliares =====
try:
    import docx  # python-docx
except Exception:
    docx = None

SUPPORTED_DOCS = {".pdf", ".docx", ".txt", ".csv", ".xlsx"}
SUPPORTED_IMAGES = {".png", ".jpg", ".jpeg", ".webp", ".tiff", ".tif"}


def _ext(name: str) -> str:
    i = name.rfind(".")
    return name[i:].lower() if i != -1 else ""


def _sha1_8(b: bytes) -> str:
    return hashlib.sha1(b).hexdigest()[:8]


def _ocr_image_bytes(img_bytes: bytes, lang: str = "spa+eng") -> str:
    if not _OCR_OK:
        return ""
    im = Image.open(io.BytesIO(img_bytes)).convert("RGB")
    return (pytesseract.image_to_string(im, lang=lang) or "").strip()


# ---------------- PDF ----------------
def _from_pdf(b: bytes, lang: str = "spa+eng") -> str:
    """
    Extrae texto de PDF. Si una página no tiene texto (PDF escaneado), hace OCR de la página.
    """
    try:
        doc = fitz.open(stream=b, filetype="pdf")
    except Exception:
        return ""

    partes = []
    for page in doc:
        txt = page.get_text("text") or ""
        if not txt.strip():
            # Render x2 para mejorar OCR
            pix = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
            ocr_txt = _ocr_image_bytes(pix.tobytes("png"), lang=lang)
            txt = ocr_txt or ""
        if txt.strip():
            partes.append(txt.strip())
    doc.close()

    return "\n".join(partes).strip()


# ---------------- DOCX ----------------
def _from_docx(b: bytes) -> str:
    if docx is None:
        return ""
    try:
        d = docx.Document(io.BytesIO(b))

        partes = []
        # Párrafos
        partes.extend(
            [p.text.strip() for p in d.paragraphs if p.text and p.text.strip()]
        )
        # Tablas: fila por fila
        for t in d.tables:
            for row in t.rows:
                celdas = []
                for cell in row.cells:
                    txt = " ".join(cell.text.split())
                    if txt:
                        celdas.append(txt)
                if celdas:
                    partes.append(" | ".join(celdas))

        return "\n".join(partes).strip()
    except Exception:
        return ""


# ---------------- TXT ----------------
def _from_txt(b: bytes) -> str:
    try:
        return b.decode("utf-8", errors="ignore").strip()
    except Exception:
        return b.decode("latin-1", errors="ignore").strip()


# ---------------- CSV ----------------
def _from_csv(b: bytes) -> str:
    out = []
    reader = csv.reader(io.StringIO(b.decode("utf-8", errors="ignore")))
    for i, row in enumerate(reader):
        out.append(", ".join(row))
        if i >= 2000:
            out.append("... (truncado)")
            break
    return "\n".join(out)


# ---------------- XLSX ----------------
def _from_xlsx(b: bytes) -> str:
    try:
        with io.BytesIO(b) as bio:
            sheets = pd.read_excel(bio, sheet_name=None)
        chunks = []
        for name, df in sheets.items():
            chunks.append(f"--- Hoja: {name} ---")
            chunks.append(df.to_csv(index=False))
        txt = "\n".join(chunks)
        return txt[:300_000] + ("... (truncado)" if len(txt) > 300_000 else "")
    except Exception:
        return ""


# ---------------- Imagen (OCR) ----------------
def _from_image(b: bytes, lang: str = "spa+eng") -> str:
    if not _OCR_OK:
        return ""
    try:
        img = Image.open(io.BytesIO(b)).convert("RGB")
        return (pytesseract.image_to_string(img, lang=lang) or "").strip()
    except Exception:
        return ""


# ================= API pública =================
def extract_attachment(name: str, content: bytes) -> Tuple[str, Dict]:
    """
    Devuelve (texto_extraído, metadatos)
    metadatos = { filename, ext, size_bytes, sha1_8, chars }
    """
    ext = _ext(name)
    text = ""

    if ext in SUPPORTED_DOCS:
        if ext == ".pdf":
            text = _from_pdf(content)
        elif ext == ".docx":
            text = _from_docx(content)
        elif ext == ".txt":
            text = _from_txt(content)
        elif ext == ".csv":
            text = _from_csv(content)
        elif ext == ".xlsx":
            text = _from_xlsx(content)
    elif ext in SUPPORTED_IMAGES:
        text = _from_image(content)

    meta = {
        "filename": name,
        "ext": ext,
        "size_bytes": len(content),
        "sha1_8": _sha1_8(content),
        "chars": len(text),
    }
    return text, meta


def consolidate_attachments(
    files: List[Tuple[str, bytes]], max_chars: int = 200_000
) -> Tuple[str, List[Dict]]:
    """
    Concatena el texto de múltiples adjuntos con encabezados por fuente y
    corta a max_chars.
    """
    parts: List[str] = []
    metas: List[Dict] = []
    total = 0

    for name, content in files:
        text, meta = extract_attachment(name, content)
        metas.append(meta)  # guardamos meta aunque no haya texto

        if not text:
            continue

        block = f"\n\n### Fuente: {name} ({meta['sha1_8']})\n{text}"
        if total + len(block) > max_chars:
            block = block[: max(0, max_chars - total)] + "\n... (truncado)"
            parts.append(block)
            break

        parts.append(block)
        total += len(block)

    return "".join(parts).strip(), metas


def ocr_diagnostics() -> Dict:
    return {
        "OCR_OK": _OCR_OK,
        "tesseract_cmd": getattr(pytesseract.pytesseract, "tesseract_cmd", None) if _OCR_OK else None,
        "engine": "PyMuPDF + Tesseract",
    }
